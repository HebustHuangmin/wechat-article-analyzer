import os
import shutil
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import re

XLSX_PATH = '结果.xlsx'
OUT_DIR = '送审签'
TEMPLATE_PATH = '模板.docx'


def clean_output_dir():
    """清空输出目录"""
    if os.path.exists(OUT_DIR):
        shutil.rmtree(OUT_DIR)
        print(f"✓ 已清空目录: {OUT_DIR}")
    os.makedirs(OUT_DIR, exist_ok=True)


def apply_font(paragraph):
    """将段落中的run字体设置为仿宋_GB2312，小四号"""
    for run in paragraph.runs:
        font = run.font
        font.name = '仿宋_GB2312'
        font.size = Pt(12)
        try:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        except Exception:
            pass


def apply_type_font(paragraph):
    """将类型段落字体设置为楷体_GB2312，三号"""
    for run in paragraph.runs:
        font = run.font
        font.name = '楷体_GB2312'
        font.size = Pt(16)
        try:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体_GB2312')
        except Exception:
            pass


def replace_text_in_runs(paragraph, old_text, new_text, keep_suffix=False, apply_default_font=True):
    """在段落的runs中替换文本（保留格式）
    
    keep_suffix: 为True时，保留占位符后面的原始文本（用于日期保留时间）
    apply_default_font: 替换后是否应用默认仿宋字体；对类型字段传 False
    """
    if old_text not in paragraph.text:
        return False

    full_text = paragraph.text

    # 如果需要保留占位符之后的文本（如日期后面的时间），只替换占位符本身
    if keep_suffix:
        paragraph.text = full_text.replace(old_text, str(new_text), 1)
        if apply_default_font:
            apply_font(paragraph)
        return True

    # 如果占位符在行首且后面带示例/选项，则直接替换成实际值，去掉示例文字
    if full_text.strip().startswith(old_text) and len(full_text.strip()) > len(old_text):
        paragraph.text = str(new_text)
        if apply_default_font:
            apply_font(paragraph)
        return True

    # 清空原有runs并按原样替换占位符
    for run in paragraph.runs:
        run.text = ""

    new_paragraph_text = full_text.replace(old_text, new_text, 1)
    paragraph.text = new_paragraph_text

    if apply_default_font:
        apply_font(paragraph)
    return True


def format_date(value):
    """将日期格式化为 'YYYY年MM月DD日'，失败则返回原值"""
    try:
        dt = pd.to_datetime(value)
        return dt.strftime('%Y年%m月%d日')
    except Exception:
        return value


def format_type(value):
    """将类型渲染为候选项列表，选中的加对勾，其余为空框；“其他”无方框并带下划线"""
    options = ["视频", "文字", "图片", "海报", "其他"]
    if value is None:
        selected = set()
    else:
        parts = [p.strip() for p in re.split('[，,]', str(value)) if p.strip()]
        selected = set(parts)

    rendered = []
    for opt in options:
        if opt == '其他':
            rendered.append("其他________")
            continue
        mark = "☑" if opt in selected else "□"
        rendered.append(f"{mark}{opt}")
    return "  ".join(rendered)


def fill_template(template_path, context):
    """填充模板文档中的{{}}占位符
    
    Args:
        template_path: 模板文件路径
        context: 包含数据的字典
    
    Returns:
        填充后的Document对象
    """
    doc = Document(template_path)
    
    # 在段落中替换文本
    for paragraph in doc.paragraphs:
        type_replaced = False
        for key, value in context.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in paragraph.text:
                # 日期特殊处理：格式化且保留后缀（如时间）
                if key == '日期':
                    replace_text_in_runs(paragraph, placeholder, format_date(value), keep_suffix=True)
                # 类型前加对勾方框，使用楷体三号（跳过默认字体）
                elif key == '类型':
                    replace_text_in_runs(paragraph, placeholder, format_type(value), apply_default_font=False)
                    type_replaced = True
                else:
                    replace_text_in_runs(paragraph, placeholder, str(value))
        if type_replaced:
            apply_type_font(paragraph)
        else:
            apply_font(paragraph)
    
    # 在表格中替换文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    type_replaced = False
                    for key, value in context.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in paragraph.text:
                            if key == '日期':
                                replace_text_in_runs(paragraph, placeholder, format_date(value), keep_suffix=True)
                            elif key == '类型':
                                replace_text_in_runs(paragraph, placeholder, format_type(value), apply_default_font=False)
                                type_replaced = True
                            else:
                                replace_text_in_runs(paragraph, placeholder, str(value))
                    if type_replaced:
                        apply_type_font(paragraph)
                    else:
                        apply_font(paragraph)
    
    return doc


def generate_word_docs(xlsx_path=XLSX_PATH):
    """从Excel结果文件使用模板生成Word文档
    
    Args:
        xlsx_path: Excel文件路径，列标题应与模板中的{{}}占位符对应
    """
    if not os.path.exists(xlsx_path):
        print(f"✗ 找不到输入文件: {xlsx_path}")
        return
    
    if not os.path.exists(TEMPLATE_PATH):
        print(f"✗ 找不到模板文件: {TEMPLATE_PATH}")
        return
    
    # 清空输出目录
    clean_output_dir()
    
    # 读取Excel文件
    try:
        df = pd.read_excel(xlsx_path, dtype=str)
        print(f"✓ 已读取 {xlsx_path}，共 {len(df)} 条记录")
    except Exception as e:
        print(f"✗ 读取Excel文件失败: {e}")
        return
    
    # 生成Word文档
    success_count = 0
    for idx, row in df.iterrows():
        try:
            # 获取日期和序号用于文件名
            date = str(row.get('日期', '')).replace('/', '-').replace('\\', '-')
            seq = str(row.get('序号', '')).strip()
            fname = f"{date}_{seq}.docx" if date else f"{seq}.docx"
            fname = fname.replace(' ', '_')
            out_path = os.path.join(OUT_DIR, fname)
            
            # 将行数据转换为字典，处理NaN值
            context = {}
            for col in df.columns:
                value = row.get(col, '')
                # 将NaN和空值转换为空字符串
                if pd.isna(value) or value == 'nan':
                    context[col] = ''
                else:
                    context[col] = str(value).strip()
            
            # 使用模板填充
            doc = fill_template(TEMPLATE_PATH, context)
            doc.save(out_path)
            
            print(f"✓ 生成: {out_path}")
            success_count += 1
        except Exception as e:
            print(f"✗ 处理第 {idx+1} 行失败 (序号: {row.get('序号', 'N/A')}): {e}")
    
    print(f"\n✓ 完成！成功生成 {success_count}/{len(df)} 个文件")


def main():
    """当直接运行此脚本时调用"""
    generate_word_docs()


if __name__ == '__main__':
    main()

